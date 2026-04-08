#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
story:export - 导出小说内容

将章节内容导出为 txt 或 docx 格式。
"""

import sys
import re
import json
import argparse
from pathlib import Path
from typing import List, Tuple, Optional
from datetime import datetime


class Colors:
    HEADER = '\033[95m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    DIM = '\033[2m'


def c(text: str, color: str) -> str:
    return f"{color}{text}{Colors.ENDC}"


# ============================================================================
# 工具函数
# ============================================================================

def find_project_root():
    """查找项目根目录"""
    cwd = Path.cwd()
    current = cwd
    for _ in range(10):
        if (current / 'story.json').exists() or (current / 'story.yml').exists():
            return current
        parent = current.parent
        if parent == current:
            break
        current = parent
    return None


def load_config(root):
    """加载配置"""
    config_path = root / 'story.json'
    if not config_path.exists():
        config_path = root / 'story.yml'

    if config_path.suffix == '.json':
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    else:
        import yaml
        with open(config_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)


def get_chapter_volume(chapter_num: int, chapters_per: int = 30) -> int:
    """根据章节号计算卷号"""
    return ((chapter_num - 1) // chapters_per) + 1


def parse_export_range(range_str: str, chapters_per: int = 30) -> Tuple[int, int]:
    """
    解析导出范围字符串。

    支持格式：
    - "5"        -> (5, 5)
    - "3-5"      -> (3, 5)
    - "3~5"      -> (3, 5)
    """
    range_str = range_str.strip().lower()

    if '-' in range_str:
        parts = range_str.split('-')
    elif '~' in range_str:
        parts = range_str.split('~')
    else:
        try:
            num = int(range_str)
            return num, num
        except ValueError:
            return 1, 1

    if len(parts) == 2:
        try:
            start = int(parts[0].strip())
            end = int(parts[1].strip())
            return start, end
        except ValueError:
            return 1, 1

    return 1, 1


def get_chapter_path(root: Path, chapter_num: int, chapters_per: int = 30) -> Path:
    """获取章节文件路径"""
    volume_num = get_chapter_volume(chapter_num, chapters_per)
    return root / 'CONTENT' / f'volume-{volume_num}' / f'chapter-{chapter_num:03d}.md'


def strip_frontmatter(content: str) -> str:
    """去除 Markdown 的 frontmatter"""
    lines = content.split('\n')
    result_lines = []
    in_frontmatter = False

    for line in lines:
        if line.strip() == '---':
            in_frontmatter = not in_frontmatter
            continue
        if in_frontmatter:
            continue
        # 跳过章节标题行（# 标题）
        if line.startswith('# '):
            continue
        result_lines.append(line)

    return '\n'.join(result_lines).strip()


def merge_chapters(chapter_contents: List[Tuple[int, str]]) -> str:
    """合并多个章节内容"""
    merged = []
    for chapter_num, content in chapter_contents:
        merged.append(f'\n{"=" * 50}\n')
        merged.append(f'第{chapter_num}章\n')
        merged.append(f'{"=" * 50}\n\n')
        merged.append(content)
        merged.append('\n\n')

    return ''.join(merged)


def count_words(text: str) -> int:
    """统计字数（中文+英文单词）"""
    # 中文字符
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    # 英文单词
    english_words = len(re.findall(r'[a-zA-Z]+', text))
    # 其他字符（标点等）
    return chinese_chars + english_words


# ============================================================================
# 导出功能
# ============================================================================

def export_to_txt(root: Path, chapters: List[int], chapters_per: int,
                  config: dict, output_path: Path) -> str:
    """导出为纯文本格式"""
    content_dir = root / 'CONTENT'
    chapter_contents = []

    for chapter_num in chapters:
        chapter_path = get_chapter_path(root, chapter_num, chapters_per)

        if not chapter_path.exists():
            print(f"  {c(f'[警告] 第{chapter_num}章不存在，跳过', Colors.YELLOW)}")
            continue

        # 读取并处理内容
        raw_content = chapter_path.read_text(encoding='utf-8')
        clean_content = strip_frontmatter(raw_content)

        if clean_content:
            chapter_contents.append((chapter_num, clean_content))

    # 合并内容
    merged = merge_chapters(chapter_contents)

    # 添加头部信息
    title = config.get('meta', {}).get('title', '未命名小说')
    header = f"{title}\n"
    header += f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
    header += f"共 {len(chapter_contents)} 章，约 {count_words(merged)} 字\n"
    header += "=" * 50 + "\n\n"

    full_content = header + merged

    # 写入文件
    output_path.write_text(full_content, encoding='utf-8')

    return full_content


def export_to_docx(root: Path, chapters: List[int], chapters_per: int,
                   config: dict, output_path: Path) -> bool:
    """导出为 docx 格式"""
    content_dir = root / 'CONTENT'
    chapter_contents = []

    for chapter_num in chapters:
        chapter_path = get_chapter_path(root, chapter_num, chapters_per)

        if not chapter_path.exists():
            print(f"  {c(f'[警告] 第{chapter_num}章不存在，跳过', Colors.YELLOW)}")
            continue

        raw_content = chapter_path.read_text(encoding='utf-8')
        clean_content = strip_frontmatter(raw_content)

        if clean_content:
            chapter_contents.append((chapter_num, clean_content))

    if not chapter_contents:
        print(f"  {c('[ERROR] 没有可导出的章节内容', Colors.RED)}")
        return False

    # 尝试使用 python-docx
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print(f"  {c('[ERROR] 导出 docx 需要安装 python-docx', Colors.RED)}")
        print(f"  {c('请运行: pip install python-docx', Colors.DIM)}")
        print(f"  {c('或使用 --format txt 导出纯文本', Colors.DIM)}")
        return False

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(12)

    # 添加标题
    title = config.get('meta', {}).get('title', '未命名小说')
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加导出信息
    info_para = doc.add_paragraph()
    info_para.add_run(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    info_para.add_run(f"共 {len(chapter_contents)} 章")
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # 添加章节内容
    for chapter_num, content in chapter_contents:
        # 章节标题
        chapter_title = doc.add_heading(f'第{chapter_num}章', 1)

        # 将 Markdown 转换为段落
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # 检测标题级别
            if line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)
            else:
                # 普通段落
                # 处理 Markdown 格式
                line = line.replace('**', '')
                para = doc.add_paragraph(line)

        doc.add_page_break()

    # 保存文档
    doc.save(str(output_path))
    return True


def ensure_export_dir(root: Path) -> Path:
    """确保导出目录存在"""
    export_dir = root / 'EXPORT'
    export_dir.mkdir(exist_ok=True)
    return export_dir


# ============================================================================
# 主函数
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='导出小说内容 - 将章节导出为 txt 或 docx 格式',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  story:export                  # 导出全部章节
  story:export 1-10            # 导出第1-10章
  story:export 5               # 导出第5章
  story:export 1-10 --format docx   # 导出为 Word 文档
  story:export --volume 1      # 导出第1卷全部章节
  story:export -o my-novel.txt # 指定输出文件名

提示：
  导出的文件保存在 EXPORT/ 目录下。
  docx 格式需要安装 python-docx：pip install python-docx
        """
    )
    parser.add_argument('range', nargs='?', help='导出范围（如 1-10 或 5）')
    parser.add_argument('--format', '-f', choices=['txt', 'docx'],
                        default='txt', help='导出格式（默认 txt）')
    parser.add_argument('--volume', '-v', type=int,
                        help='指定卷号（导出该卷全部章节）')
    parser.add_argument('--output', '-o', type=str,
                        help='输出文件名（默认自动生成）')

    args = parser.parse_args()

    root = find_project_root()
    if not root:
        print(f"  {c('[ERROR] 未找到项目目录，请先运行 story:init', Colors.RED)}")
        sys.exit(1)

    config = load_config(root)
    chapters_per = config.get('structure', {}).get('chapters_per_volume', 30)
    total_volumes = config.get('structure', {}).get('volumes', 1)

    # 确定导出范围
    if args.volume:
        # 指定卷
        vol_num = args.volume
        if vol_num < 1 or vol_num > total_volumes:
            print(f"  {c(f'[ERROR] 卷号 {vol_num} 超出范围（1-{total_volumes}）', Colors.RED)}")
            sys.exit(1)

        start_ch = (vol_num - 1) * chapters_per + 1
        end_ch = vol_num * chapters_per
        chapters = list(range(start_ch, end_ch + 1))
        default_name = f"volume-{vol_num}"
        print(f"\n{c(f'📤 导出第{vol_num}卷', Colors.CYAN)}")

    elif args.range:
        start, end = parse_export_range(args.range, chapters_per)
        if start > end:
            print(f"  {c('[ERROR] 起始章节大于结束章节', Colors.RED)}")
            sys.exit(1)
        chapters = list(range(start, end + 1))
        default_name = f"ch{start}-{end}"
        print(f"\n{c(f'📤 导出第{start}-{end}章', Colors.CYAN)}")

    else:
        # 导出全部
        total_chapters = total_volumes * chapters_per
        chapters = list(range(1, total_chapters + 1))
        default_name = "full"
        print(f"\n{c('📤 导出全部章节', Colors.CYAN)}")

    if not chapters:
        print(f"  {c('[ERROR] 没有可导出的章节', Colors.RED)}")
        sys.exit(1)

    print(f"  {c('章节数:', Colors.DIM)} {len(chapters)} 章")

    # 确定输出路径
    export_dir = ensure_export_dir(root)

    if args.output:
        # 用户指定文件名
        if not args.output.endswith(f'.{args.format}'):
            args.output += f'.{args.format}'
        output_path = export_dir / args.output
    else:
        # 自动生成文件名
        title = config.get('meta', {}).get('title', 'novel')
        # 清理文件名中的非法字符
        title = re.sub(r'[<>:"/\\|?*]', '', title)
        output_path = export_dir / f"{title}-{default_name}.{args.format}"

    # 导出
    if args.format == 'docx':
        success = export_to_docx(root, chapters, chapters_per, config, output_path)
    else:
        content = export_to_txt(root, chapters, chapters_per, config, output_path)
        success = True
        word_count = count_words(content)

    if success:
        print(f"  {c('输出文件:', Colors.DIM)} {output_path.relative_to(root)}")
        if args.format == 'txt':
            print(f"  {c('总字数:', Colors.DIM)} 约 {word_count} 字")
        print()
        print(c(f'  ✅ 导出完成！', Colors.GREEN))
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
